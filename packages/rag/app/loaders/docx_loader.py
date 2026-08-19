from __future__ import annotations

import io
import zipfile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .base import DocumentLoader

# DOCX is a zip archive; zipfile will happily decompress an arbitrarily large
# member with no size ceiling of its own. A malicious file well under the
# upload size limit can still decompress to gigabytes ("zip bomb"). Reject
# anything whose *uncompressed* content would exceed this before python-docx
# ever unpacks it. 500MB is generous for any legitimate document (a real DOCX
# rarely exceeds a few MB uncompressed) while still blocking bomb-style ratios.
_MAX_UNCOMPRESSED_BYTES = 500 * 1024 * 1024


class DocxLoader(DocumentLoader):
    """Extracts paragraphs, headings, and table cell text using python-docx.
    Headings are prefixed with '#' markers (matching Markdown-style heading
    syntax) so the chunker's heading-boundary detection (app/chunking.py) works
    uniformly across DOCX and Markdown sources.
    """

    def extract_text(self, data: bytes) -> str:
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as archive:
                total_uncompressed = sum(info.file_size for info in archive.infolist())
        except zipfile.BadZipFile as exc:
            raise ValueError(f"could not read DOCX: {exc}") from exc
        if total_uncompressed > _MAX_UNCOMPRESSED_BYTES:
            raise ValueError("DOCX contents are too large to process")

        try:
            document = Document(io.BytesIO(data))
        except (PackageNotFoundError, Exception) as exc:
            raise ValueError(f"could not read DOCX: {exc}") from exc

        parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if "heading" in style_name or "title" in style_name:
                parts.append(f"# {text}")
            else:
                parts.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        if not parts:
            raise ValueError("no extractable text found in DOCX")

        return "\n\n".join(parts)
