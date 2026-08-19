import io
import zipfile

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter
from pypdf.generic import ContentStream, DictionaryObject, NameObject

from app.loaders import get_loader
from app.loaders.docx_loader import DocxLoader
from app.loaders.pdf_loader import PdfLoader
from app.loaders.text_loader import MarkdownLoader, TxtLoader


def _make_pdf_bytes(text: str | None) -> bytes:
    """Builds a minimal single-page PDF with a real text-drawing content
    stream (no reportlab dependency needed) — or a blank page with no text
    operator at all if `text` is None, for the "no extractable text" case.
    """
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)

    if text:
        escaped = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
        content_bytes = f"BT /F1 12 Tf 72 700 Td ({escaped}) Tj ET".encode("latin-1")
        content_stream = ContentStream(None, writer)
        content_stream.set_data(content_bytes)
        page.replace_contents(content_stream)

        font_dict = DictionaryObject()
        font_dict[NameObject("/Type")] = NameObject("/Font")
        font_dict[NameObject("/Subtype")] = NameObject("/Type1")
        font_dict[NameObject("/BaseFont")] = NameObject("/Helvetica")
        font_ref = writer._add_object(font_dict)

        resources = DictionaryObject()
        fonts = DictionaryObject()
        fonts[NameObject("/F1")] = font_ref
        resources[NameObject("/Font")] = fonts
        page[NameObject("/Resources")] = resources

    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _make_docx_bytes(paragraphs: list[str], heading: str | None = None) -> bytes:
    doc = DocxDocument()
    if heading:
        doc.add_heading(heading, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_txt_loader_extracts_plain_text():
    loader = TxtLoader()
    text = loader.extract_text("Hello, this is a resume.".encode("utf-8"))
    assert "Hello, this is a resume." in text


def test_txt_loader_rejects_empty_file():
    loader = TxtLoader()
    with pytest.raises(ValueError):
        loader.extract_text(b"")


def test_txt_loader_falls_back_to_latin1_on_bad_utf8():
    loader = TxtLoader()
    data = "café".encode("latin-1")
    text = loader.extract_text(data)
    assert "caf" in text


def test_markdown_loader_preserves_heading_markers():
    loader = MarkdownLoader()
    text = loader.extract_text(b"# Project\n\nBuilt a RAG system.")
    assert "# Project" in text
    assert "Built a RAG system." in text


def test_docx_loader_extracts_paragraphs_and_headings():
    loader = get_loader(".docx")
    data = _make_docx_bytes(["I built a RAG system using FastAPI."], heading="Project Summary")
    text = loader.extract_text(data)
    assert "Project Summary" in text
    assert "I built a RAG system using FastAPI." in text
    # Heading should carry the '#' marker so the chunker recognizes it uniformly.
    assert "# Project Summary" in text


def test_docx_loader_rejects_corrupt_file():
    loader = get_loader(".docx")
    with pytest.raises(ValueError):
        loader.extract_text(b"not a real docx file")


def test_docx_loader_rejects_empty_document():
    loader = get_loader(".docx")
    data = _make_docx_bytes([])
    with pytest.raises(ValueError):
        loader.extract_text(data)


def test_pdf_loader_extracts_real_text():
    loader = get_loader(".pdf")
    data = _make_pdf_bytes("Hello RAG world")
    text = loader.extract_text(data)
    assert "Hello RAG world" in text


def test_pdf_loader_rejects_corrupt_file():
    loader = get_loader(".pdf")
    with pytest.raises(ValueError):
        loader.extract_text(b"not a real pdf file")


def test_pdf_loader_rejects_blank_page_pdf_with_no_text():
    # A structurally valid PDF with no text content should be rejected as
    # having no extractable text, not silently produce an empty document.
    loader = get_loader(".pdf")
    data = _make_pdf_bytes(None)
    with pytest.raises(ValueError):
        loader.extract_text(data)


def test_get_loader_raises_for_unknown_extension():
    with pytest.raises(ValueError):
        get_loader(".xyz")


def test_docx_loader_rejects_zip_bomb(monkeypatch):
    """A DOCX crafted to decompress far past the sane-document ceiling must
    be rejected before python-docx ever unpacks it, not after memory has
    already been exhausted. Security regression for the zip-bomb finding."""
    import app.loaders.docx_loader as docx_loader_module

    monkeypatch.setattr(docx_loader_module, "_MAX_UNCOMPRESSED_BYTES", 100)

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as archive:
        # Highly compressible content well past the (monkeypatched) 100-byte
        # uncompressed ceiling, achievable with a tiny compressed payload.
        archive.writestr("word/document.xml", b"A" * 10_000)

    with pytest.raises(ValueError, match="too large"):
        DocxLoader().extract_text(buf.getvalue())


def test_pdf_loader_rejects_too_many_pages(monkeypatch):
    """A PDF within the byte-size limit but with a pathological page count
    must be rejected rather than forcing unbounded synchronous extraction
    work. Security regression for the page-count-DoS finding."""
    import app.loaders.pdf_loader as pdf_loader_module

    monkeypatch.setattr(pdf_loader_module, "_MAX_PAGES", 2)

    writer = PdfWriter()
    for _ in range(3):
        writer.add_blank_page(width=100, height=100)
    buf = io.BytesIO()
    writer.write(buf)

    with pytest.raises(ValueError, match="too many pages"):
        PdfLoader().extract_text(buf.getvalue())
